"""
ingest.py -- читает .docx из docs/ и загружает в ChromaDB.
Читает параграфы И таблицы. FAQ-документы -- собирает пары Вопрос+Ответ из таблиц.
"""

import os
import re
import uuid
from docx import Document
from openai import OpenAI
import chromadb
from dotenv import load_dotenv

load_dotenv()

# ─── НАСТРОЙКИ ───────────────────────────────────────────────────────────────
DOCS_FOLDER   = "./docs"
COLLECTION    = "callcenter-docs"
CHROMA_FOLDER = "./chroma_db"
CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 200
EMBED_MODEL   = "text-embedding-3-small"
# ─────────────────────────────────────────────────────────────────────────────

openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
chroma_client = chromadb.PersistentClient(path=CHROMA_FOLDER)


def read_txt_file(path: str) -> str:
    """Читает текстовый файл (.txt / .md)."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def read_docx_paragraphs(path: str) -> str:
    """Читает только параграфы (стандартный текст)."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_docx_tables(path: str) -> list[str]:
    """Читает таблицы и возвращает список строк (каждая строка = один чанк)."""
    doc = Document(path)
    rows = []
    for table in doc.tables:
        # Проверяем заголовки первой строки
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]

        # Определяем структуру по заголовкам
        has_question = any('вопрос' in h or 'вопросы' in h for h in header_cells)
        has_answer   = any('ответ' in h or 'ответы' in h for h in header_cells)
        has_solution = any('решение' in h or 'способ решения' in h for h in header_cells)

        if has_question and (has_answer or has_solution):
            # FAQ-таблица: собираем пары Вопрос+Ответ
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3:
                    q_text = cells[1]  # колонка Вопрос
                    a_text = cells[2]  # колонка Ответ/Решение
                    if q_text and a_text:
                        rows.append(f"Вопрос: {q_text}\nОтвет: {a_text}")
        else:
            # Обычная таблица: каждая строка = один чанк
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
    return rows


def chunk_text(text: str) -> list[str]:
    """Механическое разбиение по 500 символов."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return [c for c in chunks if c]


def get_embedding(text: str) -> list[float]:
    if len(text) > 20000:
        text = text[:20000]
    resp = openai_client.embeddings.create(input=text, model=EMBED_MODEL)
    return resp.data[0].embedding


def main():
    docx_files = [f for f in os.listdir(DOCS_FOLDER) if f.endswith(".docx") and not f.startswith("~")]
    txt_files  = [f for f in os.listdir(DOCS_FOLDER) if f.endswith((".txt", ".md"))]
    if not docx_files:
        print("[!] Нет .docx файлов в папке docs/")
        return

    all_files = docx_files + txt_files
    print(f"[INFO] Найдено файлов: {len(all_files)} (docx: {len(docx_files)}, txt/md: {len(txt_files)})")

    try:
        chroma_client.delete_collection(COLLECTION)
        print("[OK] Старая коллекция удалена.")
    except Exception:
        pass

    collection = chroma_client.create_collection(
        name=COLLECTION,
        metadata={"hnsw:space": "cosine"}
    )
    print(f"[OK] Коллекция '{COLLECTION}' создана.")

    all_ids = []
    all_embeddings = []
    all_documents = []
    all_metadatas = []

    for filename in all_files:
        filepath = os.path.join(DOCS_FOLDER, filename)
        print(f"  Читаю: {filename}")

        chunks = []

        if filename.endswith((".txt", ".md")):
            text = read_txt_file(filepath)
            if text:
                chunks.extend(chunk_text(text))
        else:
            # 1. Параграфы
            para_text = read_docx_paragraphs(filepath)
            if para_text:
                chunks.extend(chunk_text(para_text))

            # 2. Таблицы (FAQ)
            table_rows = read_docx_tables(filepath)
            if table_rows:
                print(f"    -> {len(table_rows)} строк из таблиц (FAQ)")
                chunks.extend(table_rows)

        print(f"    -> Всего {len(chunks)} блоков")

        for chunk in chunks:
            emb = get_embedding(chunk)
            all_ids.append(str(uuid.uuid4()))
            all_embeddings.append(emb)
            all_documents.append(chunk)
            all_metadatas.append({"source": filename})

    batch_size = 100
    for i in range(0, len(all_ids), batch_size):
        collection.add(
            ids=all_ids[i:i+batch_size],
            embeddings=all_embeddings[i:i+batch_size],
            documents=all_documents[i:i+batch_size],
            metadatas=all_metadatas[i:i+batch_size],
        )
        print(f"  [INFO] Загружено {min(i+batch_size, len(all_ids))}/{len(all_ids)}...")

    print(f"\n[DONE] Готово! Загружено {len(all_ids)} блоков в ChromaDB.")
    print("Скопируй chroma_db/ в backend/chroma_db/ и redeploy.")


if __name__ == "__main__":
    main()
